from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
import re

def save_report_to_word(report_text: str, filename: str):
    lines = re.sub(r"[#*]+", "", report_text).splitlines()
    doc = Document()

    # 默认正文样式：宋体 12pt
    normal_style = doc.styles['Normal']
    normal_font = normal_style.font
    normal_font.name = '宋体'
    normal_font.size = Pt(12)
    normal_style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    for i, line in enumerate(lines):
        line = line.strip()
        if not line:
            continue
        # 第一行作为大标题
        if i == 0:
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = '黑体'
            run.font.size = Pt(16)
            run.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        # 一级标题（以“X、”开头）
        elif re.match(r"^[一二三四五六七八九十]、", line):
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = '黑体'
            run.font.size = Pt(14)
            run.bold = True
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        else:
            doc.add_paragraph(line)

    doc.save(filename)
